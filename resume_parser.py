from pathlib import Path
from io import BytesIO
import re
import json

import fitz
from docx import Document
import pandas as pd

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from PIL import Image
except Exception:
    Image = None


SKILL_SET = {
    "python", "java", "javascript", "typescript", "sql", "aws", "azure", "gcp",
    "docker", "kubernetes", "pandas", "numpy", "scikit-learn", "tensorflow",
    "pytorch", "django", "flask", "fastapi", "react", "node.js", "html", "css",
    "git", "linux", "spark", "airflow", "power bi", "tableau", "excel",
    "postgresql", "mysql", "mongodb", "rest", "graphql", "machine learning",
    "nlp", "data analysis"
}

SECTION_HEADINGS = [
    "summary", "professional summary", "profile", "objective",
    "experience", "work experience", "professional experience", "employment history",
    "skills", "technical skills", "education", "projects", "certifications",
    "achievements"
]

EXPERIENCE_MARKERS = [
    "experience", "work experience", "professional experience", "employment history"
]


def save_uploaded_file(uploaded_file) -> str:
    output_dir = Path("temp_uploads")
    output_dir.mkdir(exist_ok=True)
    file_path = output_dir / uploaded_file.name
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return str(file_path)


def extract_text_from_pdf(pdf_path: str) -> tuple[str, bool]:
    doc = fitz.open(pdf_path)
    parts = []
    ocr_used = False

    for page in doc:
        text = page.get_text("text").strip()
        if text:
            parts.append(text)
            continue

        if pytesseract is not None and Image is not None:
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                img = Image.open(BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    parts.append(ocr_text)
                    ocr_used = True
            except Exception:
                pass

    doc.close()
    return "\n".join(parts), ocr_used


def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_into_sections(text: str) -> dict:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    sections = {}
    current = "header"
    sections[current] = []

    normalized_headings = {h.lower(): h for h in SECTION_HEADINGS}

    for line in lines:
        normalized = re.sub(r"[:\-–—]+$", "", line.lower().strip()).strip()
        if normalized in normalized_headings:
            current = normalized_headings[normalized]
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


def extract_skills(text: str) -> list[str]:
    lower = text.lower()
    found = []
    for skill in SKILL_SET:
        if re.search(rf"\b{re.escape(skill)}\b", lower):
            found.append(skill)
    return sorted(set(found))


def extract_experience_section(sections: dict) -> str:
    for key in EXPERIENCE_MARKERS:
        if key in sections:
            return sections[key]
    return ""


def split_experience_blocks(exp_text: str) -> list[dict]:
    if not exp_text:
        return []

    lines = [ln.strip() for ln in exp_text.splitlines() if ln.strip()]
    blocks = []
    current_title = "Experience"
    current_lines = []

    company_or_role = re.compile(
        r"(?i)\b(inc|llc|corp|corporation|company|co\.|ltd|solutions|systems|technologies|group|studio|labs|consulting|services)\b"
    )
    date_pattern = re.compile(
        r"(?i)\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}\b|\b\d{4}\b\s*[-–]\s*(?:present|\d{4})\b"
    )

    for line in lines:
        short_line = len(line) < 130
        looks_like_header = short_line and (company_or_role.search(line) or date_pattern.search(line))

        if looks_like_header:
            if current_lines:
                blocks.append({
                    "title": current_title,
                    "content": "\n".join(current_lines).strip()
                })
            current_title = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        blocks.append({
            "title": current_title,
            "content": "\n".join(current_lines).strip()
        })

    return blocks[:10]


def summarize_experience(exp_text: str) -> list[str]:
    if not exp_text:
        return []

    lines = [ln.strip() for ln in exp_text.splitlines() if ln.strip()]
    bullets = []

    for line in lines:
        cleaned = re.sub(r"^[•*\-\u2022]+\s*", "", line).strip()
        if len(cleaned) > 20:
            bullets.append(cleaned)

    return bullets[:12]


def count_words(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def parse_resume_file(file_path: str) -> dict:
    suffix = Path(file_path).suffix.lower()
    ocr_used = False

    if suffix == ".pdf":
        raw_text, ocr_used = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

    cleaned = clean_text(raw_text)
    sections = split_into_sections(cleaned)
    experience_text = extract_experience_section(sections)
    experience_summary = summarize_experience(experience_text)
    experience_blocks = split_experience_blocks(experience_text)
    skills = extract_skills(cleaned)

    return {
        "raw_text": cleaned,
        "sections": sections,
        "skills": skills,
        "experience_summary": experience_summary,
        "experience_blocks": experience_blocks,
        "ocr_used": ocr_used,
        "stats": {
            "skills_count": len(skills),
            "sections_count": len(sections),
            "experience_lines": len(experience_summary),
            "experience_blocks_count": len(experience_blocks),
            "word_count": count_words(cleaned),
        },
    }


def results_to_json(result: dict) -> str:
    return json.dumps(result, indent=2, ensure_ascii=False)


def results_to_csv(result: dict) -> str:
    rows = []

    for skill in result.get("skills", []):
        rows.append({"type": "skill", "value": skill, "details": ""})

    for item in result.get("experience_summary", []):
        rows.append({"type": "experience_summary", "value": item, "details": ""})

    for block in result.get("experience_blocks", []):
        rows.append({
            "type": "experience_block",
            "value": block.get("title", ""),
            "details": block.get("content", "")
        })

    for section_name, section_text in result.get("sections", {}).items():
        rows.append({
            "type": "section",
            "value": section_name,
            "details": section_text
        })

    stats = result.get("stats", {})
    for k, v in stats.items():
        rows.append({"type": "stat", "value": k, "details": str(v)})

    rows.append({"type": "ocr_used", "value": str(result.get("ocr_used", False)), "details": ""})

    df = pd.DataFrame(rows)
    return df.to_csv(index=False)

